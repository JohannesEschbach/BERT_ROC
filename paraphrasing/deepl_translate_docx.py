import requests
from io import BytesIO
from typing import List, Iterable, Generator
import tempfile
import docx
import csv
import os
import time
from uuid import UUID
from zipfile import BadZipfile
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException
from tbselenium.tbdriver import TorBrowserDriver

from utils import Buckets


TBB_PATH = "~/.local/share/torbrowser/tbb/x86_64/tor-browser_en-US"
TBB_PATH = os.path.realpath(os.path.expanduser(TBB_PATH))
TIMEOUT_DEFAULT = 10
TIMEOUT_TRANSLATE = 180


def _get_session():
    session = requests.session()
    # Tor uses the 9050 port as the default socks port
    session.proxies = {'http':  'socks5://127.0.0.1:9050',
                       'https': 'socks5://127.0.0.1:9050'}
    return session


def _write_docx(docx_file: BytesIO, rows: Iterable[List[str]]):
    doc = docx.Document()
    table = doc.add_table(rows=2, cols=len(rows[0]))
    for row in rows:
        doc_cells = table.add_row().cells
        for doc_cell, input_cell in zip(doc_cells, row):
            doc_cell.text = input_cell

    doc.add_page_break()
    doc.save(docx_file)


def translate_docx(driver: WebDriver, filename: str, lang: str) -> str:
    def _f(driver: WebDriver, selector: str, by: str = By.CSS_SELECTOR) \
     -> WebElement:
        def _find(driver: WebDriver) -> WebElement:
            return driver.find_element(by, selector)
        try:
            WebDriverWait(driver, TIMEOUT_DEFAULT).until(_find)
        except TimeoutException as e:
            raise e
        return _find(driver)

    def _find_download_link(d: WebDriver) -> WebElement:
        return d.find_element(By.LINK_TEXT, 'Download again')

    for key, value in {
        "browser.download.folderList": 2,
        "browser.download.manager.showWhenStarting": False,
        "browser.helperApps.neverAsk.saveToDisk":
            "application/vnd.openxmlformats-officedocument."
            + "wordprocessingml.document"
    }.items():
        driver.profile.set_preference(key, value)
    time.sleep(10)
    driver.get("https://www.deepl.com/translator")
    try:
        _f(driver, 'button.dl_cookieBanner--buttonSelected').click()
        _f(driver, 'button[dl-test="doctrans-tabs-switch-docs"]').click()
        _f(driver, '#file-upload_input').send_keys(filename)
        _f(driver, 'button[dl-test="doctrans-upload-lang-item"]'
                   + f'[dl-lang="{lang}"]').click()
        WebDriverWait(driver,
                      TIMEOUT_TRANSLATE).until(_find_download_link)
    except TimeoutException:
        print("Restarting due to timeout")
        return
    return _find_download_link(driver).get_attribute("href")


def translate(lang: str, rows: Iterable[List[str]], bucket_size: int = 250) \
 -> Generator[List[str], None, None]:
    for bucket in Buckets(rows, bucket_size):
        with tempfile.NamedTemporaryFile(suffix=".docx") as temp_docx_in:
            _write_docx(temp_docx_in, bucket)
            while True:
                os.environ['MOZ_HEADLESS'] = '1'
                with TorBrowserDriver(TBB_PATH) as driver:
                    href = translate_docx(driver, temp_docx_in.name, lang)
                    if not href:
                        continue
                session = _get_session()
                r = session.get(href)
                try:
                    document = docx.api.Document(BytesIO(r.content))
                    break
                except BadZipfile:
                    continue
            table = document.tables[0]
            for out_row in table.rows:
                cells = [cell.text for cell in out_row.cells]
                if len(cells) and all(cells):
                    yield cells


def paraphrase(src: str, languages: List[str], dir: str):
    dir = os.path.realpath(dir)
    if src.endswith(".csv"):
        src = src[:-4]
    current_src = os.path.join(dir, f"{src}.csv")
    for i, lang in enumerate(languages):
        if not os.path.isfile(current_src):
            raise FileNotFoundError(current_src)
        filename = f"{src}.{'_'.join(languages[:i+1])}.csv"
        current_dest = os.path.join(dir, filename)
        if os.path.isfile(current_dest):
            print("completing existing file:", current_dest)
            with open(current_src) as csv_in:
                csv_in_reader = csv.reader(csv_in)
                missing_rows = []
                with open(current_dest) as csv_dest:
                    csv_dest_reader = csv.reader(csv_dest)
                    for row in csv_in_reader:
                        uuid = row[0]
                        try:
                            UUID(uuid)
                        except ValueError:
                            continue
                        for row_dest in csv_dest_reader:
                            if uuid == row[0]:
                                break
                        else:
                            missing_rows.append(row)
                print(f"{len(missing_rows)} missing rows")
                rows = translate(lang, missing_rows)
                with open(current_dest, "a") as csv_out:
                    csv_writer = csv.writer(csv_out)
                    for row in rows:
                        csv_writer.writerow(row)
        else:
            print("Translating", current_src, "to", current_dest)
            with open(current_src) as csv_in:
                csv_in_reader = csv.reader(csv_in)
                rows = translate(lang, csv_in_reader)
                with open(current_dest, "a") as csv_out:
                    csv_writer = csv.writer(csv_out)
                    for row in rows:
                        csv_writer.writerow(row)
        current_src = current_dest


languages_available = [
    'bg-BG', 'cs-CS', 'da-DA', 'de-DE', 'el-EL', 'en-GB', 'en-US', 'es-ES',
    'et-ET', 'fi-FI', 'hu-HU', 'it-IT', 'ja-JA', 'lv-LV', 'nl-NL', 'pl-PL',
    'pt-BR', 'pt-PT', 'ro-RO', 'ru-RU', 'sk-SK', 'sl-SL', 'sv-SV', 'zh-ZH'
]

if __name__ == "__main__":
    lang_input = "en"
    lang_target = "en-GB"
    lang_pipelines = [[middle, lang_target] for middle in languages_available
                      if not middle.startswith(lang_input)]
    datasets = [
        "cloze_test",
        # "cloze_test_nolabel",
    ]
    dir = os.path.realpath(os.path.join(os.getcwd(), "data"))
    for dataset in datasets:
        for languages in lang_pipelines:
            paraphrase(dataset, languages, dir)
