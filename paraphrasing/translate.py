import os
import time
from uuid import UUID
from typing import List, Iterable, Generator, Tuple
from abc import ABC, abstractmethod

from utils import Buckets, is_parsable

# used for GoogleTranslator only:
import googletrans
from httpcore._exceptions import ConnectTimeout, ReadTimeout

# used for DeepLTranslator only:
import requests
import tempfile
from io import BytesIO
import docx
from zipfile import BadZipfile
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.remote.webdriver import WebDriver
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException, \
                                       ElementClickInterceptedException, \
                                       NoSuchElementException
from tbselenium.tbdriver import TorBrowserDriver


TBB_PATH = "~/.local/share/torbrowser/tbb/x86_64/tor-browser_en-US"
TBB_PATH = os.path.realpath(os.path.expanduser(TBB_PATH))
TBB_PROXY = "socks5://127.0.0.1:9050"
TIMEOUT_DEFAULT = 10
TIMEOUT_DEEPL_TRANSLATE = 210

if __name__ == "__main__":
    wd = os.path.realpath(os.path.join(os.getcwd(), "data"))


class Translator(ABC):
    dest_languages: List[str]

    @abstractmethod
    def translate(self, lang: str, rows: Iterable[List[str]]) \
     -> Generator[List[str], None, None]:
        pass

    @staticmethod
    def _translatable(text: str) -> bool:
        if is_parsable(text, UUID):
            return False
        if is_parsable(text, int):
            return False
        return len(text) > 0

    @classmethod
    def _split_translatable(self, rows: Iterable[Iterable[str]]) \
     -> Tuple[List[List[bool]], List[str], List[str]]:
        mask, translate_input, residue = [], [], []
        for row in rows:
            mask.append([])
            for cell in row:
                mask[-1].append(self._translatable(cell))
                if mask[-1][-1]:
                    translate_input.append(cell)
                else:
                    residue.append(cell)
        return mask, translate_input, residue

    @staticmethod
    def _merge_translated(mask: List[List[bool]], translated: List[str],
                          residue: List[str]) \
     -> Generator[List[str], None, None]:
        for row in mask:
            yield [translated.pop(0) if is_translated else residue.pop(0)
                   for is_translated in row]


class GoogleTranslator(Translator):
    bucket_size: int
    dest_languages = sorted(list(googletrans.LANGUAGES.keys()))

    def __init__(self, bucket_size: int = 5):
        self.bucket_size = bucket_size

    def translate(self, lang: str, rows: Iterable[List[str]]) \
     -> Generator[List[str], None, None]:
        for j, bucket in enumerate(Buckets(rows, self.bucket_size)):
            print(f"translating bucket {j} to {lang}")
            mask, translate_input, residue = self._split_translatable(bucket)
            while True:
                start = time.time()
                translator = googletrans.Translator()
                try:
                    translated = translator.translate(translate_input,
                                                      dest=lang[:2])
                    break
                except Exception as e:
                    duration = time.time() - start
                    print(f"{e.__class__.__name__} after {int(duration):3d}s")
                    time.sleep(10)
                    print("Retrying")
            print(f"translated {len(bucket)} stories in "
                  + f"{int(time.time()-start):3d}s")
            translated = [item.text for item in translated]
            merged = self._merge_translated(mask, translated, residue)
            equal_counter = 0
            for original_row, translated_row in zip(bucket, merged):
                if original_row == translated_row:
                    equal_counter += 1
                else:
                    yield translated_row
            if equal_counter:
                print(f"{equal_counter:2d} stories returned untranslated")


class DeepLTranslator(Translator):
    bucket_size: int
    headless: bool
    dest_languages = [
        'bg-BG', 'cs-CS', 'da-DA', 'de-DE', 'el-EL', 'en-GB', 'en-US', 'es-ES',
        'et-ET', 'fi-FI', 'hu-HU', 'it-IT', 'ja-JA', 'lv-LV', 'nl-NL', 'pl-PL',
        'pt-BR', 'pt-PT', 'ro-RO', 'ru-RU', 'sk-SK', 'sl-SL', 'sv-SV', 'zh-ZH'
    ]

    def __init__(self, bucket_size: int = 250, headless: bool = True):
        self.bucket_size = bucket_size
        self.headless = headless

    def translate(self, lang: str, rows: Iterable[List[str]]) \
     -> Generator[List[str], None, None]:
        for i, bucket in enumerate(Buckets(rows, self.bucket_size)):
            mask, translate_input, residue = self._split_translatable(bucket)
            with tempfile.NamedTemporaryFile(suffix=".docx") as temp_docx_in:
                self._write_docx(temp_docx_in, [[c] for c in translate_input])
                while True:
                    if self.headless:
                        os.environ['MOZ_HEADLESS'] = '1'
                    with TorBrowserDriver(TBB_PATH) as driver:
                        try:
                            href = self._translate_docx(driver,
                                                        temp_docx_in.name,
                                                        lang)
                            print(href)
                        except Exception as e:
                            _debug_screenshot(driver, e)
                            raise e
                            continue
                        if not href:
                            print("no href")
                            continue
                    session = TorSession()
                    r = session.get(href)
                    try:
                        document = docx.api.Document(BytesIO(r.content))
                        break
                    except BadZipfile:
                        continue
                table = document.tables[0]
                translated = [cell.text
                              for out_row in table.rows[1:]
                              for cell in out_row.cells]
                print(translated)
                if not len(translated) or not all(translated):
                    print(f"nothing translated from bucket {i}!")
                    continue
                yield from self._merge_translated(mask, translated, residue)

    @staticmethod
    def _write_docx(docx_file: BytesIO, rows: Iterable[List[str]]):
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=len(rows[0]))
        for row in rows:
            doc_cells = table.add_row().cells
            for doc_cell, input_cell in zip(doc_cells, row):
                doc_cell.text = input_cell

        doc.add_page_break()
        doc.save(docx_file)

    @staticmethod
    def _translate_docx(driver: WebDriver, filename: str, lang: str) -> str:
        def _f(driver: WebDriver, selector: str, by: str = By.CSS_SELECTOR) \
         -> WebElement:
            def _find(driver: WebDriver) -> WebElement:
                return driver.find_element(by, selector)
            try:
                WebDriverWait(driver, TIMEOUT_DEFAULT).until(_find)
            except TimeoutException as e:
                raise e
            return _find(driver)

        def _find_download_link(d: WebDriver) -> WebElement:
            return d.find_element(By.LINK_TEXT, 'Download again')

        for key, value in {
            "browser.download.folderList": 2,
            "browser.download.manager.showWhenStarting": False,
            "browser.helperApps.neverAsk.saveToDisk":
                "application/vnd.openxmlformats-officedocument."
                + "wordprocessingml.document"
        }.items():
            driver.profile.set_preference(key, value)
        time.sleep(2)
        try:
            driver.get("https://www.deepl.com/translator")
            # _f(driver, 'button.dl_cookieBanner--buttonSelected').click()
            _f(driver, 'button[dl-test="doctrans-tabs-switch-docs"]').click()
            _f(driver, '#file-upload_input').send_keys(filename)
            _f(driver, 'button[dl-test="doctrans-upload-lang-item"]'
                       + f'[dl-lang="{lang}"]').send_keys(Keys.RETURN)
            WebDriverWait(driver,
                          TIMEOUT_DEEPL_TRANSLATE).until(_find_download_link)
        except TimeoutException:
            print("Restarting due to timeout")
        except ElementClickInterceptedException as e:
            _debug_screenshot(driver, e)
            return
        try:
            return _find_download_link(driver).get_attribute("href")
        except NoSuchElementException as e:
            _debug_screenshot(driver, e)


class TorSession(requests.sessions.Session):
    def __init__(self):
        super().__init__()
        # Tor uses the 9050 port as the default socks port
        self.proxies = {'http': TBB_PROXY, 'https': TBB_PROXY}


def _debug_screenshot(driver: WebDriver, e: Exception,
                      wd: str = wd if __name__ == "__main__" else "",
                      do: bool = __name__ == "__main__"):
    if do:
        filename = f"{e.__class__.__name__}_{int(time.time())}.png"
        driver.save_screenshot(os.path.join(wd, filename))
